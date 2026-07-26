"""文档解析器：统一产出带结构的 Block 列表（文本/标题/表格），供切分器消费。

设计要点：
- 解析结果保留 heading 层级与页码，供结构感知切分与引用溯源
- PDF 快速通道用 PyMuPDF，含页眉页脚启发式剔除与断行修复
- quality 通道（MinerU）预留接口，一期未实现时回退 fast
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field


@dataclass
class Block:
    """文档结构化块。kind: heading / text / table。"""
    kind: str
    text: str
    level: int = 0          # heading 层级，1 为最顶层
    page: int | None = None
    meta: dict = field(default_factory=dict)


@dataclass
class ParserOutput:
    blocks: list[Block]
    source_name: str


# ---------------------------------------------------------------- 清洗工具

_CN_END_PUNCT = "。！？；：…）】》”’—"
_EN_END_PUNCT = ".!?;:)]}\"'"


def clean_text(text: str) -> str:
    """基础清洗：控制字符、空白规范化。"""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def fix_broken_lines(text: str) -> str:
    """修复 PDF 抽取的硬换行：行尾非标点且下一行非标题样式时拼接。"""
    lines = text.split("\n")
    out: list[str] = []
    for line in lines:
        s = line.rstrip()
        if not s:
            out.append("")
            continue
        if out and out[-1] and out[-1][-1] not in _CN_END_PUNCT + _EN_END_PUNCT:
            # 中文直接拼接，英文补空格
            joiner = "" if _is_cjk(out[-1][-1]) else " "
            out[-1] = out[-1] + joiner + s.lstrip()
        else:
            out.append(s)
    return "\n".join(out)


def _is_cjk(ch: str) -> bool:
    return "一" <= ch <= "鿿"


def _looks_like_header_footer(line: str, page_no: int, total: int) -> bool:
    """启发式识别页眉页脚：短行 + 页码特征 + 重复模式。"""
    s = line.strip()
    if not s:
        return True
    if len(s) <= 4 and re.fullmatch(r"[-—–\s]*\d+[-—–\s]*", s):  # 页码 " - 42 - "
        return True
    if re.fullmatch(rf"第\s*\d+\s*页(\s*共\s*\d+\s*页)?", s):
        return True
    if s == str(page_no) or s == f"{page_no} / {total}":
        return True
    return False


# ---------------------------------------------------------------- PDF

class PdfParser:
    """PyMuPDF 快速通道。"""

    def parse(self, data: bytes, name: str) -> ParserOutput:
        import fitz  # PyMuPDF，延迟导入

        blocks: list[Block] = []
        with fitz.open(stream=io.BytesIO(data), filetype="pdf") as doc:
            total = doc.page_count
            for page_no, page in enumerate(doc, start=1):
                raw = page.get_text("text")
                lines = [
                    ln for ln in raw.split("\n")
                    if not _looks_like_header_footer(ln, page_no, total)
                ]
                text = fix_broken_lines(clean_text("\n".join(lines)))
                for para in re.split(r"\n{2,}", text):
                    para = para.strip()
                    if not para:
                        continue
                    heading = self._detect_heading(para)
                    if heading:
                        level, title = heading
                        blocks.append(Block("heading", title, level=level, page=page_no))
                    else:
                        blocks.append(Block("text", para, page=page_no))
        return ParserOutput(blocks=blocks, source_name=name)

    _RE_CHAPTER = re.compile(r"^(第[一二三四五六七八九十百\d]+[章节篇部分])\s*(.*)$")
    _RE_NUMBERED = re.compile(r"^(\d+(?:\.\d+){0,3})[、.\s]\s*(\S.*)$")

    def _detect_heading(self, para: str) -> tuple[int, str] | None:
        """单行且匹配章节/编号模式时判为标题。"""
        s = para.strip()
        if "\n" in s or len(s) > 60:
            return None
        m = self._RE_CHAPTER.match(s)
        if m:
            return 1, s
        m = self._RE_NUMBERED.match(s)
        if m and not s.endswith("。"):
            depth = m.group(1).count(".") + 1
            return min(depth, 4), s
        return None


# ---------------------------------------------------------------- DOCX

class DocxParser:
    def parse(self, data: bytes, name: str) -> ParserOutput:
        import docx  # python-docx，延迟导入

        blocks: list[Block] = []
        document = docx.Document(io.BytesIO(data))
        for para in document.paragraphs:
            text = clean_text(para.text)
            if not text:
                continue
            style = (para.style.name or "").lower() if para.style else ""
            if style.startswith("heading"):
                level = self._style_level(style)
                blocks.append(Block("heading", text, level=level))
            else:
                blocks.append(Block("text", text))
        # 表格整块输出为 Markdown
        for table in document.tables:
            md = self._table_to_markdown(table)
            if md:
                blocks.append(Block("table", md))
        return ParserOutput(blocks=blocks, source_name=name)

    @staticmethod
    def _style_level(style: str) -> int:
        m = re.search(r"(\d+)", style)
        return int(m.group(1)) if m else 1

    @staticmethod
    def _table_to_markdown(table) -> str:
        rows = []
        for row in table.rows:
            cells = [clean_text(c.text).replace("|", "\\|") for c in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if not rows:
            return ""
        header_sep = "|" + " --- |" * len(table.rows[0].cells)
        return "\n".join([rows[0], header_sep, *rows[1:]])


# ---------------------------------------------------------------- MD / TXT

class MarkdownParser:
    def parse(self, data: bytes, name: str) -> ParserOutput:
        text = clean_text(data.decode("utf-8", errors="replace"))
        blocks: list[Block] = []
        buf: list[str] = []

        def flush() -> None:
            joined = "\n".join(buf).strip()
            if joined:
                blocks.append(Block("text", joined))
            buf.clear()

        for line in text.split("\n"):
            m = re.match(r"^(#{1,6})\s+(.*)$", line)
            if m:
                flush()
                blocks.append(Block("heading", m.group(2).strip(), level=len(m.group(1))))
            else:
                buf.append(line)
        flush()
        return ParserOutput(blocks=blocks, source_name=name)


# ---------------------------------------------------------------- 注册表

_PARSERS = {
    "pdf": PdfParser(),
    "docx": DocxParser(),
    "doc": DocxParser(),
    "md": MarkdownParser(),
    "markdown": MarkdownParser(),
    "txt": MarkdownParser(),
}


def parse_document(data: bytes, file_type: str, name: str, parser: str = "fast") -> ParserOutput:
    """入口：按类型分发。quality 通道预留（MinerU），一期回退 fast。"""
    ft = (file_type or "").lower().lstrip(".")
    impl = _PARSERS.get(ft)
    if impl is None:
        raise ValueError(f"unsupported file type: {file_type}")
    return impl.parse(data, name)
